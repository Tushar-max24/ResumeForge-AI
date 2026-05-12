from docx import Document
import os

templates = [
    "d:/TalentVerse AI/Tushar/ResumeForge AI/Backend/templates/CV of Humres.docx",
    "d:/TalentVerse AI/Tushar/ResumeForge AI/Backend/templates/CV of HunTek.docx",
    "d:/TalentVerse AI/Tushar/ResumeForge AI/Backend/templates/CV of Totaco Template.docx"
]

for t_path in templates:
    print(f"\n--- Checking {os.path.basename(t_path)} ---")
    doc = Document(t_path)
    
    found = False
    # Check Tables
    for i, table in enumerate(doc.tables):
        if len(table.rows) > 0:
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    if "CURRICULUM" in cell.text.upper() and "VITAE" in cell.text.upper():
                        print(f"Found in Table {i}, Row {r_idx}, Cell {c_idx}: '{cell.text}'")
                        found = True
    
    # Check Paragraphs
    for i, p in enumerate(doc.paragraphs):
        if "CURRICULUM" in p.text.upper() and "VITAE" in p.text.upper():
            print(f"Found in Paragraph {i}: '{p.text}'")
            found = True
    
    if not found:
        print("NOT FOUND")
