from docx import Document

def inspect_template(path):
    doc = Document(path)
    print(f"Inspecting {path}")
    
    print("\n--- Paragraphs ---")
    for i, p in enumerate(doc.paragraphs):
        if 'CURRICULUM VITAE' in p.text.upper():
            print(f"Paragraph {i}: '{p.text}'")
            for r in p.runs:
                print(f"  Run: '{r.text}', Size: {r.font.size.pt if r.font.size else 'None'}, Bold: {r.bold}")

    print("\n--- Tables ---")
    for i, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                if 'CURRICULUM VITAE' in cell.text.upper():
                    print(f"Table {i}, Row {r_idx}, Cell {c_idx}: '{cell.text}'")
                    for p in cell.paragraphs:
                        for run in p.runs:
                            print(f"  Run: '{run.text}', Size: {run.font.size.pt if run.font.size else 'None'}, Bold: {run.bold}")

if __name__ == "__main__":
    inspect_template("d:/TalentVerse AI/Tushar/ResumeForge AI/Backend/templates/CV of Totaco Template.docx")
