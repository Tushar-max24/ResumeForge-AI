from docx import Document
import os

files = [
    "d:/TalentVerse AI/Tushar/ResumeForge AI/Backend/scratch/test_humres.docx",
    "d:/TalentVerse AI/Tushar/ResumeForge AI/Backend/scratch/test_huntek.docx",
    "d:/TalentVerse AI/Tushar/ResumeForge AI/Backend/scratch/test_totaco.docx"
]

for f_path in files:
    print(f"\n--- Checking {os.path.basename(f_path)} ---")
    doc = Document(f_path)
    
    # Check Tables
    for i, table in enumerate(doc.tables):
        if len(table.rows) > 0:
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    if "CURRICULUM VITAE" in cell.text.upper():
                        print(f"Found in Table {i}, Row {r_idx}, Cell {c_idx}")
                        for p in cell.paragraphs:
                            for r in p.runs:
                                size = r.font.size.pt if r.font.size else "Default"
                                print(f"  Run text: '{r.text}', Size: {size}")
    
    # Check Paragraphs
    for i, p in enumerate(doc.paragraphs):
        if "CURRICULUM VITAE" in p.text.upper():
            print(f"Found in Paragraph {i}")
            for r in p.runs:
                size = r.font.size.pt if r.font.size else "Default"
                print(f"  Run text: '{r.text}', Size: {size}")
