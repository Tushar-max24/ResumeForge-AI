import json
import os
import subprocess
from docx import Document

def verify_fix():
    template_path = "d:/TalentVerse AI/Tushar/ResumeForge AI/Backend/templates/CV of Totaco Template.docx"
    output_path = "d:/TalentVerse AI/Tushar/ResumeForge AI/Backend/scratch/verified_totaco.docx"
    json_path = "d:/TalentVerse AI/Tushar/ResumeForge AI/Backend/scratch/dummy_data.json"
    
    # Dummy data
    data = {
        "personal_info": {
            "name": "John Doe",
            "location": "New York, USA",
            "sector": "Technology",
            "rl_id": "RL12345"
        },
        "personal_profile": "A professional software engineer.",
        "employment_summary": [],
        "comprehensive_work_history": [],
        "education_and_skills": {}
    }
    
    with open(json_path, 'w') as f:
        json.dump(data, f)
    
    # Run the script
    print("Running fill_docx.py...")
    cmd = ["python", "d:/TalentVerse AI/Tushar/ResumeForge AI/Backend/scripts/fill_docx.py", template_path, output_path, json_path]
    result = subprocess.run(cmd, capture_output=True, text=True)
    
    if result.returncode != 0:
        print(f"Error running script: {result.stderr}")
        return
    
    print("Script finished successfully.")
    
    # Verify the font size
    doc = Document(output_path)
    found = False
    for table in doc.tables:
        if len(table.rows) > 0 and 'CURRICULUM VITAE' in table.rows[0].cells[0].text.upper():
            found = True
            print(f"Found CURRICULUM VITAE in Table 0")
            for p in table.rows[0].cells[0].paragraphs:
                for r in p.runs:
                    size = r.font.size.pt if r.font.size else "None"
                    print(f"  Run: '{r.text}', Size: {size}")
                    if size == 16.0:
                        print("  VERIFICATION SUCCESS: Font size is 16.0")
                    else:
                        print(f"  VERIFICATION FAILURE: Font size is {size}")
    
    if not found:
        print("VERIFICATION FAILURE: CURRICULUM VITAE heading not found in output document.")

if __name__ == "__main__":
    verify_fix()
