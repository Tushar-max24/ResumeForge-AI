from docx import Document

def verify_output(path):
    doc = Document(path)
    print(f"Verifying {path}\n")
    
    career_history_started = False
    for p in doc.paragraphs:
        text = p.text.strip()
        if 'CAREER HISTORY' in text.upper():
            career_history_started = True
            print(f"Found Career History section.")
            continue
            
        if career_history_started:
            if text:
                print(f"Paragraph: '{text[:100]}...'")
                # Check for bullet points and "Key responsibilities included"
                if text.startswith('•\tKey responsibilities included:'):
                    print("  [SUCCESS] Found joined responsibilities bullet.")
    # Education check
    edu_section_started = False
    edu_items = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if 'EDUCATION' in text.upper():
            edu_section_started = True
            print(f"\nFound Education section.")
            continue
            
        if edu_section_started:
            if text.startswith('•\t'):
                item_text = text.replace('•\t', '').strip()
                edu_items.append(item_text)
                print(f"  Edu Item: '{item_text}'")
            elif text and not any(k in text.upper() for k in ['SKILLS', 'AWARDS', 'LICENSE']):
                # This might be a header or other text
                pass
            elif any(k in text.upper() for k in ['SKILLS', 'AWARDS', 'LICENSE']):
                break
    
    # Verify order
    years = []
    import re
    for item in edu_items:
        match = re.search(r'\b(19|20)\d{2}\b', item)
        if match:
            years.append(int(match.group(0)))
    
    if years == sorted(years):
        print(f"  [SUCCESS] Education items are sorted ascending: {years}")
    else:
        print(f"  [FAILURE] Education items are NOT sorted ascending: {years}")

if __name__ == "__main__":
    verify_output("d:/TalentVerse AI/Tushar/ResumeForge AI/Backend/scratch/test_output.docx")
