import sys
import json
import os
from docx import Document
from docx.shared import Pt, Inches
from dateutil import parser as date_parser
from datetime import datetime

def format_date(date_str):
    if not date_str:
        return ''
    if date_str.lower() in ['present', 'current', 'ongoing', 'now']:
        return 'Present'
    try:
        dt = date_parser.parse(date_str)
        return dt.strftime('%b-%Y')
    except:
        return date_str

def get_date_obj(date_str):
    if not date_str or date_str.lower() in ['present', 'current', 'ongoing', 'now']:
        return datetime.now()
    try:
        return date_parser.parse(date_str)
    except:
        return datetime(1900, 1, 1)

def set_font(run, size=10, bold=False, italic=False):
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.all_caps = False

def format_location(location):
    """Extract city and country from full address"""
    if not location:
        return ''
    
    # Remove common address components and keep only city and country
    location = location.strip()
    
    # Split by comma and take last parts (usually city, country format)
    if ',' in location:
        parts = [p.strip() for p in location.split(',')]
        # Take last two parts (city, country) or just last part if only two
        if len(parts) >= 2:
            return ', '.join(parts[-2:])
        elif len(parts) == 1:
            return parts[0]
    
    # If no comma, try to extract city and country from address
    # Remove street numbers, apartment numbers, etc.
    words = location.split()
    city_country_words = []
    
    for word in words:
        # Skip common street/address indicators
        if any(indicator in word.lower() for indicator in ['street', 'road', 'avenue', 'lane', 'drive', 'boulevard', 'apt', 'suite', 'unit', '#']):
            continue
        # Keep words that look like city or country (usually longer words)
        if len(word) > 2 and not word.isdigit():
            city_country_words.append(word)
    
    if city_country_words:
        return ', '.join(city_country_words[-2:])  # Take last 2 meaningful words
    elif words:
        # Fallback: return last word if it looks like a city
        return words[-1]
    
    return location.title()

def fill_resume(template_path, output_path, data):
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template not found: {template_path}")
    
    doc = Document(template_path)
    
    # Extract data parts
    personal_info = data.get('personal_info', {})
    profile = data.get('personal_profile', '')
    employment_summary = data.get('employment_summary', [])
    career_history = data.get('comprehensive_work_history', []) or data.get('career_history', [])
    ed_skills = data.get('education_and_skills', {})

    # Helper to find table by first row text
    def find_table(text):
        for table in doc.tables:
            if table.rows and any(text.lower() in cell.text.lower() for cell in table.rows[0].cells):
                return table
        return None

    # Fill Personal Info
    is_huntek = "HunTek" in template_path
    is_totaco = "Totaco" in template_path
    
    info_table = None
    for table in doc.tables:
        if len(table.rows) > 0 and 'Name:' in table.rows[0].cells[0].text:
            info_table = table
            break
    
    if info_table:
        if is_huntek:
            # HunTek Format: Name + RL ID on top, Location below
            # Format name to show only first and last name
            raw_name = personal_info.get('name', '').strip()
            name_parts = raw_name.split()
            if len(name_parts) >= 2:
                name_val = f"{name_parts[0].title()} {name_parts[-1].title()}"
            else:
                name_val = raw_name.title()
            rl_id_val = personal_info.get('rl_id', '')
            loc_val = format_location(personal_info.get('location', ''))
            
            # Combine Name and RL ID in the first value cell
            name_cell = info_table.cell(0, 1)
            name_cell.text = f"{name_val} {rl_id_val}".strip()
            for p in name_cell.paragraphs:
                p.paragraph_format.alignment = 0 # LEFT
                p.paragraph_format.space_before = 0
                p.paragraph_format.space_after = 0
                for r in p.runs: set_font(r, 11)
            
            # Put Location in the second row
            if len(info_table.rows) > 1:
                loc_cell = info_table.cell(1, 1)
                loc_cell.text = loc_val
                for p in loc_cell.paragraphs:
                    p.paragraph_format.alignment = 0 # LEFT
                    p.paragraph_format.space_before = 0
                    p.paragraph_format.space_after = 0
                    for r in p.runs: set_font(r, 11)
                
                # Clear RL ID and Sector from the other cells if they exist
                if len(info_table.rows[0].cells) > 3:
                    info_table.cell(0, 3).text = ""
                    info_table.cell(1, 3).text = ""
        else:
            # Standard Totaco/Humres Format
            # Row 0: Name and Location
            for i in [0, 2]:
                for p in info_table.cell(0, i).paragraphs:
                    for r in p.runs: set_font(r, 11, True)
            
            name_cell = info_table.cell(0, 1)
            # Format name to show only first and last name
            raw_name = personal_info.get('name', '').strip()
            name_parts = raw_name.split()
            if len(name_parts) >= 2:
                formatted_name = f"{name_parts[0].title()} {name_parts[-1].title()}"
            else:
                formatted_name = raw_name.title()
            name_cell.text = formatted_name
            for p in name_cell.paragraphs:
                for r in p.runs: set_font(r, 11)
                
            loc_cell = info_table.cell(0, 3)
            loc_cell.text = format_location(personal_info.get('location', ''))
            for p in loc_cell.paragraphs:
                for r in p.runs: set_font(r, 11)
    
            # Row 1: Sector and RL ID
            if len(info_table.rows) > 1:
                for i in [0, 2]:
                    for p in info_table.cell(1, i).paragraphs:
                        for r in p.runs: set_font(r, 11, True)
                        
                sec_cell = info_table.cell(1, 1)
                sec_cell.text = personal_info.get('sector', '')
                for p in sec_cell.paragraphs:
                    for r in p.runs: set_font(r, 11)
                    
                rl_cell = info_table.cell(1, 3)
                rl_cell.text = personal_info.get('rl_id', '')
                for p in rl_cell.paragraphs:
                    for r in p.runs: set_font(r, 11)


    # Helper to find section starting point (Table or Paragraph)
    def find_section_header(keywords):
        # 1. Search Tables
        for table in doc.tables:
            if len(table.rows) > 0:
                h_text = table.rows[0].cells[0].text.upper()
                if all(k.upper() in h_text for k in keywords):
                    return table, "table"
        # 2. Search Paragraphs
        for i, p in enumerate(doc.paragraphs):
            h_text = p.text.upper()
            if all(k.upper() in h_text for k in keywords):
                return p, "paragraph"
        return None, None

    # Set "CURRICULUM VITAE" font size to 16
    cv_header, cv_type = find_section_header(['CURRICULUM', 'VITAE'])
    if cv_header:
        if cv_type == "table":
            for p in cv_header.rows[0].cells[0].paragraphs:
                for r in p.runs: set_font(r, 16, True)
        else:
            for r in cv_header.runs: set_font(r, 16, True)

    # Fill Profile
    if profile:
        header_el, type = find_section_header(['PROFILE'])
        if header_el:
            if type == "table":
                last_el = header_el._element
                # Set Main Heading Font for table header
                for p in header_el.rows[0].cells[0].paragraphs:
                    for r in p.runs: set_font(r, 16, True)
            else:
                last_el = header_el._element
                # Set Main Heading Font for paragraph
                for r in header_el.runs: set_font(r, 16, True)

            # Add spacer first
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_after = Pt(8)
            last_el.addnext(spacer._element)
            
            # Insert profile text
            p = doc.add_paragraph(profile)
            p.paragraph_format.line_spacing = 1.15
            for r in p.runs: set_font(r, 11)
            spacer._element.addnext(p._element)

    # Fill Employment Summary Table
    emp_header, _ = find_section_header(['SUMMARY', 'EMPLOYMENT'])
    if emp_header and hasattr(emp_header, 'rows'):
        for p in emp_header.rows[0].cells[0].paragraphs:
            for r in p.runs: set_font(r, 16, True)

    summary_table = None
    for table in doc.tables:
        if len(table.rows) > 0:
            h_text = table.rows[0].cells[0].text.upper()
            if 'FROM' in h_text or 'DATES' in h_text:
                summary_table = table
                break
    
    if summary_table:
        # Set Table Column Headers Font (Sub-headings)
        for cell in summary_table.rows[0].cells:
            for p in cell.paragraphs:
                for r in p.runs: set_font(r, 11, True)
        
        # Clear existing empty rows (except header)
        while len(summary_table.rows) > 1:
            row = summary_table.rows[1]
            row._element.getparent().remove(row._element)
        
        num_cols = len(summary_table.rows[0].cells)
        
        # 1. Group by Company (Case-insensitive)
        grouped = {}
        for job in employment_summary:
            company_raw = job.get('company_name', '').strip()
            if not company_raw: continue
            company_key = company_raw.lower()
            
            if company_key not in grouped:
                grouped[company_key] = {
                    'from': job.get('from', ''),
                    'to': job.get('to', ''),
                    'company_name': company_raw.title(),
                    'position': job.get('position', ''),
                }
            else:
                # Same company: merge tenure, keep latest position
                try:
                    cur_from = get_date_obj(job.get('from', ''))
                    cur_to = get_date_obj(job.get('to', ''))
                    existing_from = get_date_obj(grouped[company_key]['from'])
                    existing_to = get_date_obj(grouped[company_key]['to'])
                    
                    if cur_from < existing_from:
                        grouped[company_key]['from'] = job.get('from', '')
                    if cur_to > existing_to:
                        grouped[company_key]['to'] = job.get('to', '')
                        grouped[company_key]['position'] = job.get('position', '')
                except:
                    pass

        # 2. Sort and Limit to Top 5
        processed_summary = list(grouped.values())
        processed_summary.sort(key=lambda x: get_date_obj(x.get('from', '')), reverse=True)
        processed_summary = processed_summary[:5]

        # 3. Fill Table
        for job in processed_summary:
            row = summary_table.add_row()
            cells = row.cells
            
            job_from = format_date(job.get('from', ''))
            job_to = format_date(job.get('to', ''))
            job_company = job.get('company_name', '').title()
            job_pos = job.get('position', '')

            # Set column widths to prevent date wrapping
            try:
                summary_table.allow_autofit = False
                if num_cols >= 6: # Totaco
                    summary_table.columns[0].width = Inches(1.3) # FROM
                    summary_table.columns[2].width = Inches(1.3) # TO
                    summary_table.columns[4].width = Inches(2.2) # COMPANY
                    summary_table.columns[5].width = Inches(2.2) # POSITION
                elif num_cols >= 4: # Humres
                    summary_table.columns[0].width = Inches(1.3) # FROM
                    summary_table.columns[1].width = Inches(1.3) # TO
                    summary_table.columns[2].width = Inches(2.5) # COMPANY
                    summary_table.columns[3].width = Inches(2.5) # POSITION
            except:
                pass

            if num_cols >= 7: # Totaco style
                cells[0].text = job_from
                cells[2].text = job_to
                cells[4].text = job_company
                cells[5].text = job_pos
            elif num_cols >= 4: # Humres/Huntek style
                cells[0].text = job_from
                cells[1].text = job_to
                cells[2].text = job_company
                cells[3].text = job_pos

            # Formatting for table cells: Add padding via paragraph spacing
            for i, cell in enumerate(cells):
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(6)
                    paragraph.paragraph_format.space_after = Pt(6)
                    paragraph.paragraph_format.left_indent = 0
                    paragraph.paragraph_format.right_indent = 0
                for run in paragraph.runs:
                    set_font(run, 11)

    # Fill Career History
    career_header_el, type = find_section_header(['CAREER', 'HISTORY'])
    if career_header_el:
        if type == "table":
            last_element = career_header_el._element
            for p in career_header_el.rows[0].cells[0].paragraphs:
                for r in p.runs: set_font(r, 16, True)
        else:
            last_element = career_header_el._element
            for r in career_header_el.runs: set_font(r, 16, True)

        last_element = career_header_el._element
        # Add a spacer paragraph after the header table
        spacer_p = doc.add_paragraph()
        spacer_p.paragraph_format.space_after = Pt(6)
        last_element.addnext(spacer_p._element)
        last_element = spacer_p._element
        # Use Career History exactly as provided by AI (already sorted and complete)
        for job in career_history:
            company = job.get('company', '').strip()
            role = job.get('role', '').strip()
            period_raw = job.get('period', '').strip().replace(' to ', ' - ')
            summary = job.get('summary', '').strip()
            resps = job.get('responsibilities', [])

            # Skip entries that are clearly empty or likely misclassified References
            if not company or (not period_raw and not summary and not resps):
                continue

            # Header: COMPANY - Role (Period)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(2)
            
            company = job.get('company', '').strip().title()
            location = job.get('location', '')
            role = job.get('role', '')
            
            # Format period: "Mar-2021 - Jul-2024" -> "Mar 2021 - Jul 2024"
            if period_raw and period_raw != '()':
                if ' - ' in period_raw:
                    p_parts = period_raw.split(' - ')
                    period = f"{p_parts[0].strip().replace('-', ' ')} - {p_parts[1].strip().replace('-', ' ')}"
                else:
                    period = period_raw.replace('-', ' ')
                duration_str = f" ({period})"
            else:
                duration_str = ""
            
            # Format: Company Name – Location - Title (Duration)
            loc_str = f" – {location}" if location else ""
            run = p.add_run(f"{company}{loc_str} - {role}{duration_str}")
            set_font(run, 11, True)
            last_element.addnext(p._element)
            last_element = p._element
            
            # 1. Short Summary & Responsibilities
            summary = job.get('summary', '')
            period_val = job.get('period', '').lower()
            is_current = "current" in period_val or "present" in period_val
            resps = job.get('responsibilities', [])
            
            if is_totaco:
                # Totaco style: Everything as bulleted paragraphs
                if summary:
                    p_sum = doc.add_paragraph(style='Normal')
                    p_sum.paragraph_format.left_indent = Pt(36)
                    p_sum.paragraph_format.first_line_indent = Pt(-18)
                    p_sum.paragraph_format.space_after = Pt(6)
                    run = p_sum.add_run(f"•\t{summary}")
                    set_font(run, 11)
                    last_element.addnext(p_sum._element)
                    last_element = p_sum._element
                
                if resps:
                    p_resp = doc.add_paragraph(style='Normal')
                    p_resp.paragraph_format.left_indent = Pt(36)
                    p_resp.paragraph_format.first_line_indent = Pt(-18)
                    p_resp.paragraph_format.space_after = Pt(6)
                    
                    resp_text = ", ".join([r.strip() for r in resps if r.strip()])
                    run = p_resp.add_run(f"•\tKey responsibilities included: {resp_text}")
                    set_font(run, 11)
                    last_element.addnext(p_resp._element)
                    last_element = p_resp._element
            else:
                # Original logic for other templates
                if summary:
                    if is_current:
                        # Current Job: Summary as Paragraph (2-3 lines style)
                        p_sum = doc.add_paragraph(summary, style='Normal')
                        p_sum.paragraph_format.space_after = Pt(8)
                        p_sum.paragraph_format.line_spacing = 1.15
                        for r in p_sum.runs: set_font(r, 11)
                        last_element.addnext(p_sum._element)
                        last_element = p_sum._element
                        
                        # Add "Key responsibilities:" bold heading
                        p_key = doc.add_paragraph()
                        run = p_key.add_run("Key responsibilities:")
                        set_font(run, 11, True) # Bold
                        p_key.paragraph_format.space_after = Pt(4)
                        last_element.addnext(p_key._element)
                        last_element = p_key._element
                    else:
                        # Past Jobs: Summary as First Bullet Point
                        p_sum = doc.add_paragraph(style='Normal')
                        p_sum.paragraph_format.left_indent = Pt(36)
                        p_sum.paragraph_format.first_line_indent = Pt(-18)
                        p_sum.paragraph_format.space_after = Pt(6)
                        run = p_sum.add_run(f"•\t{summary}")
                        set_font(run, 11)
                        last_element.addnext(p_sum._element)
                        last_element = p_sum._element
                
                if resps:
                    for resp in resps:
                        if not resp.strip(): continue
                        p_resp = doc.add_paragraph(style='Normal')
                        p_resp.paragraph_format.left_indent = Pt(36)
                        p_resp.paragraph_format.first_line_indent = Pt(-18)
                        p_resp.paragraph_format.space_after = Pt(6)
                        run = p_resp.add_run(f"•\t{resp.strip()}")
                        set_font(run, 11)
                        last_element.addnext(p_resp._element)
                        last_element = p_resp._element
            
            # 3. Projects Section (Optional)
            projects = job.get('projects', [])
            if projects:
                p_proj_label = doc.add_paragraph()
                p_proj_label.paragraph_format.left_indent = Pt(18)
                p_proj_label.paragraph_format.space_before = Pt(8)
                p_proj_label.paragraph_format.space_after = Pt(4)
                run = p_proj_label.add_run("Projects")
                set_font(run, 11, True)
                last_element.addnext(p_proj_label._element)
                last_element = p_proj_label._element
                
                for proj in projects:
                    if not proj.strip(): continue
                    p_proj = doc.add_paragraph(style='Normal')
                    p_proj.paragraph_format.left_indent = Pt(36)
                    p_proj.paragraph_format.first_line_indent = Pt(-18)
                    p_proj.paragraph_format.space_after = Pt(6)
                    run = p_proj.add_run(f"•\t{proj.strip()}")
                    set_font(run, 11)
                    last_element.addnext(p_proj._element)
                    last_element = p_proj._element

            # 4. Reason for leaving Bullet
            reason = job.get('reason_for_leaving')
            if reason:
                p_reason = doc.add_paragraph(style='Normal')
                if is_totaco:
                    p_reason.paragraph_format.left_indent = Pt(36)
                    p_reason.paragraph_format.first_line_indent = Pt(-18)
                else:
                    p_reason.paragraph_format.left_indent = Pt(18)
                
                p_reason.paragraph_format.space_before = Pt(8)
                p_reason.paragraph_format.space_after = Pt(6)
                
                bullet_char = "•\t" if is_totaco else "• "
                run = p_reason.add_run(f"{bullet_char}Reason for leaving - {reason}")
                set_font(run, 11)
                last_element.addnext(p_reason._element)
                last_element = p_reason._element



            
            # Spacer
            spacer = doc.add_paragraph()
            last_element.addnext(spacer._element)
            last_element = spacer._element

    # Fill Education & Skills
    ed_header_el, type = find_section_header(['EDUCATION'])
    if not ed_header_el: # Try alternate
         ed_header_el, type = find_section_header(['SKILLS'])

    if ed_header_el:
        if type == "table":
            last_element = ed_header_el._element
            for p in ed_header_el.rows[0].cells[0].paragraphs:
                for r in p.runs: set_font(r, 16, True)
        else:
            last_element = ed_header_el._element
            for r in ed_header_el.runs: set_font(r, 16, True)

        last_element = ed_header_el._element
        
        # Add a spacer after the black header table
        spacer_p = doc.add_paragraph()
        spacer_p.paragraph_format.space_after = Pt(8)
        last_element.addnext(spacer_p._element)
        last_element = spacer_p._element
        
        # Cleanup: Remove existing redundant headers from the template to avoid duplication
        redundant_headers = [
            "Education / Qualifications", 
            "Education/Qualifications",
            "Technical Skills", 
            "Skills/Technical Skills",
            "Skills",
            "Certifications and Training", 
            "Certifications",
            "Training",
            "Awards",
            "License"
        ]
        
        # We iterate in reverse to safely delete paragraphs
        for p in reversed(doc.paragraphs):
            p_text = p.text.strip().lower()
            
            # 1. Remove empty paragraphs at the very end of the document
            if not p_text and p == doc.paragraphs[-1]:
                p._element.getparent().remove(p._element)
                continue

            # 2. Remove redundant headers
            for header in redundant_headers:
                if header.lower() in p_text and len(p_text) < len(header) + 3:
                    p._element.getparent().remove(p._element)
                    break

        # Helper to add a section with bullets
        def add_bullet_section(title, items, is_edu=False):
            nonlocal last_element
            if not items: return
            
            # Header
            p_label = doc.add_paragraph()
            p_label.paragraph_format.space_before = Pt(12)
            p_label.paragraph_format.space_after = Pt(4)
            run = p_label.add_run(title)
            set_font(run, 11, True)
            last_element.addnext(p_label._element)
            last_element = p_label._element
            
            # Bullets
            if is_edu:
                def get_yr(s):
                    import re
                    try: 
                        match = re.search(r'\b(19|20)\d{2}\b', s)
                        if match:
                            return int(match.group(0))
                        return 9999
                    except: return 9999
                items.sort(key=get_yr)

            for item in items:
                if not item.strip(): continue
                p_item = doc.add_paragraph(style='Normal')
                p_item.paragraph_format.left_indent = Pt(36)
                p_item.paragraph_format.first_line_indent = Pt(-18)
                p_item.paragraph_format.space_after = Pt(6)
                run = p_item.add_run(f"•\t{item.strip()}")
                set_font(run, 11)
                last_element.addnext(p_item._element)
                last_element = p_item._element

        # Render sections in order
        if is_totaco:
            # Consolidate Education, Training, and Certifications for Totaco
            all_edu_raw = (ed_skills.get('qualifications', []) or []) + \
                          (ed_skills.get('training', []) or []) + \
                          (ed_skills.get('certifications', []) or [])
            
            
            all_edu = []
            import re
            
            def extract_and_format_education(item):
                """Extract year from education item and format it properly"""
                if not item or not item.strip():
                    return None
                
                item = item.strip()
                
                # Try to find any year in the text (more flexible patterns)
                year_patterns = [
                    r'\b(19|20)\d{2}\b',  # Standard years like 1990-2029
                    r'\b\d{4}\b'          # Any 4-digit number (catch-all)
                ]
                
                year_found = None
                for pattern in year_patterns:
                    match = re.search(pattern, item)
                    if match:
                        year_found = match.group(0)
                        break
                
                if year_found:
                    # Remove year from original position and clean up
                    item_without_year = item.replace(year_found, '').replace('  ', ' ').strip(' -(),')
                    # Format as "Year - Degree - Institution"
                    result = f"{year_found} - {item_without_year}"
                    return result
                else:
                    # No year found, return as is
                    return item
            
            # Process all education items
            processed_items = []
            for item in all_edu_raw:
                formatted_item = extract_and_format_education(item)
                if formatted_item:
                    processed_items.append(formatted_item)
            
            # Sort by year (items without year go to the end)
            def get_sort_key(item):
                # Extract year from the beginning of the item
                year_match = re.match(r'^(\d{4})', item)
                if year_match:
                    try:
                        year_int = int(year_match.group(1))
                        return year_int
                    except:
                        return 9999  # Put at end if conversion fails
                return 9999  # Put items without year at end
            
            
            processed_items.sort(key=get_sort_key)
            
            
            # Final formatting for display
            for item in processed_items:
                parts = item.split(' - ', 2)  # Split into max 3 parts
                if len(parts) >= 3:
                    year = parts[0].strip()
                    degree = parts[1].strip()
                    institution = parts[2].strip().title()
                    result = f"{year} - {degree} - {institution}"
                    all_edu.append(result)
                elif len(parts) == 2:
                    # Could be "Year - Degree" or "Degree - Institution"
                    year_match = re.match(r'^\d{4}', parts[0])
                    if year_match:
                        # Year - Degree format
                        result = f"{parts[0].strip()} - {parts[1].strip()}"
                        all_edu.append(result)
                    else:
                        # Degree - Institution format (no year)
                        degree = parts[0].strip()
                        institution = parts[1].strip().title()
                        result = f"{degree} - {institution}"
                        all_edu.append(result)
                else:
                    # Single item
                    result = item.strip()
                    all_edu.append(result)
            

            add_bullet_section("Education/Qualification", all_edu, is_edu=True)
        else:
            add_bullet_section("Education/Qualification", ed_skills.get('qualifications', []), is_edu=True)
            add_bullet_section("Training", ed_skills.get('training', []))
            add_bullet_section("Certifications", ed_skills.get('certifications', []))

        add_bullet_section("Skills/Technical Skills", ed_skills.get('technical_skills', []))
        add_bullet_section("Awards", ed_skills.get('awards', []))
        add_bullet_section("License", ed_skills.get('license', []))



    doc.save(output_path)

if __name__ == "__main__":
    if len(sys.argv) < 4:
        print("Usage: python fill_docx.py <template_path> <output_path> <json_file_path>")
        sys.exit(1)
    
    t_path = sys.argv[1]
    o_path = sys.argv[2]
    json_file = sys.argv[3]
    try:
        with open(json_file, 'r') as f:
            j_data = json.load(f)
        fill_resume(t_path, o_path, j_data)
        print(f"Successfully generated docx at {o_path}")
    except Exception as e:
        print(f"Error: {str(e)}")
        sys.exit(1)

