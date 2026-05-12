
import docx
import sys
import os

def extract_text_from_docx(file_path):
    doc = docx.Document(file_path)
    full_text = []
    # Extract from paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
    # Extract from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                full_text.append(" | ".join(row_text))
    return '\n'.join(full_text)

files = [
    r"D:\Tushar\CV of Humres.docx",
    r"D:\Tushar\CV of HunTek.docx",
    r"D:\Tushar\CV of Totaco Template.docx"
]

for file_path in files:
    print(f"--- CONTENT OF {os.path.basename(file_path)} ---")
    try:
        text = extract_text_from_docx(file_path)
        print(text)
    except Exception as e:
        print(f"Error reading {file_path}: {e}")
    print("\n" + "="*50 + "\n")
